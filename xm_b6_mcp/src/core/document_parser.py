"""
Document Parser for SDK Documentation
======================================

Parse Word, PDF, and other documents into structured format for indexing.

Uses MarkItDown as primary parser (supports Word, PDF, HTML, etc.).
Falls back to python-docx and PyPDF2 if MarkItDown is unavailable.

Author: B6x MCP Server Team
Version: 0.1.0
"""

import os
import logging
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, NamedTuple
from dataclasses import dataclass, asdict
from datetime import datetime
import re

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Structured representation of a parsed document.

    Attributes:
        doc_id: Unique document identifier (hash of path + content)
        file_path: Original file path
        file_name: Original file name
        doc_type: Document type ('word', 'pdf', 'excel', 'markdown')
        domain: Classified domain ('domain1', 'domain2', 'domain3', 'domain4', 'domain5')
        title: Document title (extracted from filename or content)
        markdown_text: Full document content as Markdown
        sections: List of section headers found
        metadata: File metadata (size, modified_time, page_count)
        parse_method: Method used to parse ('markitdown', 'docx', 'pypdf2')
    """
    doc_id: str
    file_path: str
    file_name: str
    doc_type: str
    domain: str
    title: str
    markdown_text: str
    sections: List[str]
    metadata: Dict[str, any]
    parse_method: str

    def to_dict(self) -> Dict:
        """Convert to dictionary for JSON serialization."""
        return asdict(self)


class DocumentClassifier:
    """
    Classify documents into 5 knowledge domains based on content and filename.

    Domains:
    - Domain 1 (Hardware): Registers, electrical specs, pins, Layout, PCB
    - Domain 2 (Drivers): API, drivers, initialization, configuration
    - Domain 3 (BLE): BLE, Bluetooth, GATT, Profile, Mesh
    - Domain 4 (Applications): Examples, applications, FAQ, common issues
    - Domain 5 (Production): Mass production, programming, JFlash, certification
    """

    # Precise filename-to-domain mapping (overrides keyword-based classification)
    # Format: filename_pattern -> domain
    FILENAME_DOMAIN_MAP = {
        # Domain 1: Hardware & Registers
        "B6x_DataSheet": "domain1",
        "Layout设计错误": "domain1",
        "硬件设计指导": "domain1",
        "开发板使用说明": "domain1",
        "B6x_BLE芯片使用指南": "domain1",  # Chip-level BLE hardware
        "musbfsfc": "domain1",             # USB specifications (musbfsfc_pg.pdf, musbfsfc_ps.pdf)
        "Datasheet_Reference": "domain1",   # Datasheet_Reference/Datasheet.pdf
        "Developmentboard": "domain1",    # Development board schematics (B6_32D_20230729.pdf)

        # Domain 2: SOC Drivers
        "B6x_api": "domain2",
        "设计指导": "domain2",                # API design guide

        # Domain 3: BLE Stack
        "B6x_BLE-Sram空间分配": "domain3",
        "B6x_BLE-功耗参考": "domain3",
        "B6x BLE兼容性列表": "domain3",

        # Domain 4: Applications
        "B6x常见应用解答": "domain4",
        "JFlash配置说明": "domain4",

        # Domain 5: Production & Toolchain
        "无线认证指导": "domain5",
        "认证指导手册": "domain5",
        "bxISP_Programmer": "domain5",
        "ISP Tool量产说明文档": "domain5",
        "B6X_RF_LabTest": "domain5",        # RF test tool
        "RF_LabTest": "domain5",             # RF test tool
        "LabTest": "domain5",                # Lab test tool
        "Programmer": "domain5",             # ISP Programmer (bxISP_Programmer_User_Guide.pdf)
    }

    # Domain detection keywords (fallback when no filename match)
    DOMAIN_KEYWORDS = {
        "domain1": [
            "寄存器", "电气参数", "引脚", "pin", "register",
            "layout", "pcb", "flash", "sram", "memory map",
            "datasheet", "io_map", "flash-map", "硬件", "hardware",
            "schematic", "典型layout", "设计错误"
        ],
        "domain2": [
            "api", "驱动", "driver", "初始化", "initialization",
            "配置", "configuration", "外设", "peripheral",
            "gpio", "uart", "spi", "i2c", "timer", "pwm",
            "adc", "dac", "dma"
        ],
        "domain3": [
            "ble", "蓝牙", "bluetooth", "gatt", "profile",
            "mesh", "gap", "sm", "l2cap", "att",
            "广播", "advertising", "连接", "connection",
            "功耗", "power", "兼容性", "compatibility",
            "sram空间", "ble芯片", "协议栈"
        ],
        "domain4": [
            "示例", "example", "应用", "application",
            "faq", "常见问题", "common issue", "解答",
            "answer", "demo", "sample", "tutorial"
        ],
        "domain5": [
            "量产", "production", "烧录", "programming",
            "jflash", "isp", "tool", "工具",
            "认证", "certification", "fcc", "ce", "srrc",
            "测试", "test", "校准", "calibration",
            "packaging", "包装", "出货"
        ]
    }

    @classmethod
    def classify(cls, content: str, filename: str) -> str:
        """
        Classify a document into a domain based on content and filename.

        First checks precise filename mapping, then falls back to keyword matching.

        Args:
            content: Document content (markdown or text)
            filename: Document filename

        Returns:
            Domain identifier ('domain1' through 'domain5')
        """
        # Step 1: Check precise filename mapping first (highest priority)
        for pattern, domain in cls.FILENAME_DOMAIN_MAP.items():
            if pattern in filename:
                logger.debug(f"File '{filename}' matched pattern '{pattern}' → {domain}")
                return domain

        # Step 2: Fall back to keyword-based classification
        combined_text = (content + " " + filename).lower()

        # Score each domain
        scores = {}
        for domain, keywords in cls.DOMAIN_KEYWORDS.items():
            score = 0
            for keyword in keywords:
                # Count occurrences (case-insensitive)
                count = combined_text.count(keyword.lower())
                score += count
            scores[domain] = score

        # Find domain with highest score
        max_score = max(scores.values())
        if max_score == 0:
            # Default to domain4 if no keywords found
            return "domain4"

        # Return domain with highest score
        for domain, score in scores.items():
            if score == max_score:
                logger.debug(f"File '{filename}' keyword-matched → {domain}")
                return domain

        return "domain4"  # Fallback


class SectionExtractor:
    """Extract section headers from Markdown text."""

    # Markdown heading patterns
    HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

    @classmethod
    def extract(cls, markdown: str) -> List[str]:
        """
        Extract section headers from Markdown text.

        Args:
            markdown: Markdown text

        Returns:
            List of section headers (without # symbols)
        """
        sections = []
        for match in cls.HEADING_PATTERN.finditer(markdown):
            level = len(match.group(1))  # Number of # symbols
            title = match.group(2).strip()
            sections.append(f"[{'#' * level}] {title}")
        return sections


class DocumentParser:
    """
    Parse Word, PDF, and other documents into structured format.

    Primary method: MarkItDown (supports Word, PDF, HTML, etc.)
    Fallback methods: python-docx, PyPDF2
    """

    def __init__(self):
        """Initialize the parser with available backends."""
        self._markitdown_available = False
        self._docx_available = False
        self._pypdf2_available = False
        self._openpyxl_available = False

        # Check for MarkItDown (primary)
        try:
            from markitdown import MarkItDown
            self.MarkItDown = MarkItDown
            self._markitdown_available = True
            logger.info("MarkItDown available for document parsing")
        except ImportError:
            logger.warning("MarkItDown not available, will use fallback parsers")

        # Check for python-docx (Word fallback)
        try:
            import docx
            self.docx = docx
            self._docx_available = True
            logger.info("python-docx available for Word parsing")
        except ImportError:
            logger.warning("python-docx not available")

        # Check for PyPDF2 (PDF fallback)
        try:
            import PyPDF2
            self.PyPDF2 = PyPDF2
            self._pypdf2_available = True
            logger.info("PyPDF2 available for PDF parsing")
        except ImportError:
            logger.warning("PyPDF2 not available")

        # Check for openpyxl (Excel fallback)
        try:
            import openpyxl
            self._openpyxl_available = True
            logger.info("openpyxl available for Excel parsing")
        except ImportError:
            logger.warning("openpyxl not available")

    def _generate_doc_id(self, file_path: str, content_preview: str) -> str:
        """Generate unique document ID from path and content."""
        content_hash = hashlib.md5(
            (file_path + content_preview[:500]).encode('utf-8')
        ).hexdigest()[:16]
        return f"doc_{content_hash}"

    def _get_file_metadata(self, file_path: str) -> Dict[str, any]:
        """Get file metadata."""
        path = Path(file_path)
        stat = path.stat()

        return {
            "size_bytes": stat.st_size,
            "size_kb": round(stat.st_size / 1024, 2),
            "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "extension": path.suffix.lower()
        }

    def parse(self, file_path: str) -> Optional[ParsedDocument]:
        """
        Parse a document file (auto-detect type).

        Args:
            file_path: Path to document file

        Returns:
            ParsedDocument if successful, None if failed
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        # Route to appropriate parser
        if ext in ['.docx', '.doc']:
            return self.parse_word(file_path)
        elif ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext in ['.md', '.markdown']:
            return self.parse_markdown(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self.parse_excel(file_path)
        else:
            logger.error(f"Unsupported file type: {ext}")
            return None

    def parse_word(self, docx_path: str) -> Optional[ParsedDocument]:
        """
        Parse Word (.docx) document to Markdown.

        Args:
            docx_path: Path to Word document

        Returns:
            ParsedDocument if successful, None if failed
        """
        path = Path(docx_path)
        markdown_text = ""
        parse_method = "unknown"

        try:
            # Try MarkItDown first
            if self._markitdown_available:
                md = self.MarkItDown()
                result = md.convert(str(path))
                markdown_text = result.text_content
                parse_method = "markitdown"
                logger.info(f"Parsed {docx_path} using MarkItDown")
            elif self._docx_available:
                # Fallback to python-docx
                doc = self.docx.Document(docx_path)
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        # Detect headings
                        style_name = para.style.name if para.style else ""
                        if 'Heading' in style_name:
                            level = style_name.replace('Heading ', '')
                            text = f"{'#' * int(level)} {text}"
                        paragraphs.append(text)
                markdown_text = "\n\n".join(paragraphs)
                parse_method = "docx"
                logger.info(f"Parsed {docx_path} using python-docx")
            else:
                logger.error("No Word parser available (MarkItDown or python-docx)")
                return None

        except Exception as e:
            logger.error(f"Error parsing Word document {docx_path}: {e}")
            return None

        # Extract metadata
        metadata = self._get_file_metadata(docx_path)
        metadata["page_count"] = markdown_text.count('\n# ')  # Approximate

        # Extract sections
        sections = SectionExtractor.extract(markdown_text)

        # Classify domain
        domain = DocumentClassifier.classify(markdown_text, path.name)

        # Generate doc ID
        doc_id = self._generate_doc_id(docx_path, markdown_text)

        # Extract title
        title = path.stem.replace('_', ' ').replace('-', ' ').title()

        return ParsedDocument(
            doc_id=doc_id,
            file_path=str(path.absolute()),
            file_name=path.name,
            doc_type="word",
            domain=domain,
            title=title,
            markdown_text=markdown_text,
            sections=sections,
            metadata=metadata,
            parse_method=parse_method
        )

    def parse_pdf(self, pdf_path: str) -> Optional[ParsedDocument]:
        """
        Parse PDF document to Markdown.

        Args:
            pdf_path: Path to PDF document

        Returns:
            ParsedDocument if successful, None if failed
        """
        path = Path(pdf_path)
        markdown_text = ""
        parse_method = "unknown"

        try:
            # Try MarkItDown first
            if self._markitdown_available:
                md = self.MarkItDown()
                result = md.convert(str(path))
                markdown_text = result.text_content
                parse_method = "markitdown"
                logger.info(f"Parsed {pdf_path} using MarkItDown")
            elif self._pypdf2_available:
                # Fallback to PyPDF2
                import PyPDF2
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pages_text = []
                    for page in reader.pages:
                        text = page.extract_text()
                        pages_text.append(f"```\n{text}\n```")
                    markdown_text = "\n\n".join(pages_text)
                parse_method = "pypdf2"
                logger.info(f"Parsed {pdf_path} using PyPDF2")
            else:
                logger.error("No PDF parser available (MarkItDown or PyPDF2)")
                return None

        except Exception as e:
            logger.error(f"Error parsing PDF document {pdf_path}: {e}")
            return None

        # Extract metadata
        metadata = self._get_file_metadata(pdf_path)

        # Try to get page count
        if self._pypdf2_available:
            try:
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata["page_count"] = len(reader.pages)
            except:
                metadata["page_count"] = 0

        # Extract sections
        sections = SectionExtractor.extract(markdown_text)

        # Classify domain
        domain = DocumentClassifier.classify(markdown_text, path.name)

        # Generate doc ID
        doc_id = self._generate_doc_id(pdf_path, markdown_text)

        # Extract title
        title = path.stem.replace('_', ' ').replace('-', ' ').title()

        return ParsedDocument(
            doc_id=doc_id,
            file_path=str(path.absolute()),
            file_name=path.name,
            doc_type="pdf",
            domain=domain,
            title=title,
            markdown_text=markdown_text,
            sections=sections,
            metadata=metadata,
            parse_method=parse_method
        )

    def parse_markdown(self, md_path: str) -> Optional[ParsedDocument]:
        """
        Parse Markdown document (just read and structure).

        Args:
            md_path: Path to Markdown document

        Returns:
            ParsedDocument if successful, None if failed
        """
        path = Path(md_path)

        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
        except Exception as e:
            logger.error(f"Error reading Markdown file {md_path}: {e}")
            return None

        # Extract metadata
        metadata = self._get_file_metadata(md_path)

        # Extract sections
        sections = SectionExtractor.extract(markdown_text)

        # Classify domain
        domain = DocumentClassifier.classify(markdown_text, path.name)

        # Generate doc ID
        doc_id = self._generate_doc_id(md_path, markdown_text)

        # Extract title
        title = path.stem.replace('_', ' ').replace('-', ' ').title()

        return ParsedDocument(
            doc_id=doc_id,
            file_path=str(path.absolute()),
            file_name=path.name,
            doc_type="markdown",
            domain=domain,
            title=title,
            markdown_text=markdown_text,
            sections=sections,
            metadata=metadata,
            parse_method="native"
        )

    def parse_excel(self, xlsx_path: str) -> Optional[ParsedDocument]:
        """
        Parse Excel document to Markdown for full-text search indexing.

        Note: This is for indexing Excel content to Whoosh.
        For structured constraint extraction, use excel_parser.py instead.

        Args:
            xlsx_path: Path to Excel document

        Returns:
            ParsedDocument if successful, None if failed
        """
        path = Path(xlsx_path)
        markdown_text = ""
        parse_method = "unknown"

        try:
            # Try MarkItDown first (best for Excel → Markdown)
            if self._markitdown_available:
                md = self.MarkItDown()
                result = md.convert(str(path))
                markdown_text = result.text_content
                parse_method = "markitdown"
                logger.info(f"Parsed {xlsx_path} using MarkItDown")
            elif self._openpyxl_available:
                # Fallback: Use openpyxl to convert Excel to Markdown table
                import openpyxl
                wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)

                md_lines = [f"# {path.stem}\n"]

                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    md_lines.append(f"\n## Sheet: {sheet_name}\n")

                    # Convert to Markdown table
                    for row in sheet.iter_rows(values_only=True):
                        # Convert row to Markdown table row
                        md_row = "| " + " | ".join(
                            str(cell) if cell is not None else "" for cell in row
                        ) + " |"
                        md_lines.append(md_row)

                wb.close()
                markdown_text = "\n".join(md_lines)
                parse_method = "openpyxl"
                logger.info(f"Parsed {xlsx_path} using openpyxl")
            else:
                logger.error("No Excel parser available (MarkItDown or openpyxl)")
                return None

        except Exception as e:
            logger.error(f"Error parsing Excel document {xlsx_path}: {e}")
            return None

        # Extract metadata
        metadata = self._get_file_metadata(xlsx_path)

        # Count sheets if openpyxl is available
        if self._openpyxl_available:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(xlsx_path, read_only=True)
                metadata["sheet_count"] = len(wb.sheetnames)
                wb.close()
            except:
                metadata["sheet_count"] = 0

        # Extract sections
        sections = SectionExtractor.extract(markdown_text)

        # Classify domain
        domain = DocumentClassifier.classify(markdown_text, path.name)

        # Generate doc ID
        doc_id = self._generate_doc_id(xlsx_path, markdown_text)

        # Extract title
        title = path.stem.replace('_', ' ').replace('-', ' ').title()

        return ParsedDocument(
            doc_id=doc_id,
            file_path=str(path.absolute()),
            file_name=path.name,
            doc_type="excel",
            domain=domain,
            title=title,
            markdown_text=markdown_text,
            sections=sections,
            metadata=metadata,
            parse_method=parse_method
        )

    def parse_directory(
        self,
        directory: str,
        extensions: List[str] = None
    ) -> List[ParsedDocument]:
        """
        Parse all supported documents in a directory.

        Args:
            directory: Directory path
            extensions: File extensions to include (default: .docx, .pdf, .xlsx, .md)

        Returns:
            List of ParsedDocument objects
        """
        if extensions is None:
            extensions = ['.docx', '.doc', '.pdf', '.xlsx', '.md', '.markdown']

        path = Path(directory)
        parsed_docs = []

        for ext in extensions:
            for file_path in path.rglob(f'*{ext}'):
                logger.info(f"Parsing document: {file_path}")
                doc = self.parse(str(file_path))
                if doc:
                    parsed_docs.append(doc)

        logger.info(f"Parsed {len(parsed_docs)} documents from {directory}")
        return parsed_docs


# ============================================================================
# Convenience Functions
# ============================================================================

def parse_document(file_path: str) -> Optional[ParsedDocument]:
    """Convenience function to parse a single document."""
    parser = DocumentParser()
    return parser.parse(file_path)


def parse_documents(directory: str) -> List[ParsedDocument]:
    """Convenience function to parse all documents in a directory."""
    parser = DocumentParser()
    return parser.parse_directory(directory)
